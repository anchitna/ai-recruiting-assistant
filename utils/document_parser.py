import zipfile
import tempfile
import os
import logging
from pathlib import Path

import PyPDF2
import docx
from docx.opc.exceptions import PackageNotFoundError

logger = logging.getLogger(__name__)

class DocumentParser:
    """Utility class for extracting text from various document formats with robust error handling."""
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Extract text from a document file, automatically detecting the file type.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text as a string
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return ""
        
        file_extension = file_path.suffix.lower()
        
        if file_extension in ['.pdf']:
            return DocumentParser.extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return DocumentParser.extract_text_from_docx(file_path)
        elif file_extension in ['.txt', '.text', '.md', '.csv']:
            return DocumentParser.extract_text_from_txt(file_path)
        else:
            logger.warning(f"Unknown file extension: {file_extension}. Attempting to detect format.")
            return DocumentParser.auto_detect_and_extract(file_path)
    
    @staticmethod
    def extract_text_from_pdf(file_path: Path) -> str:
        """
        Extract text from a PDF file with robust error handling.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text
        """
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                
                if len(reader.pages) == 0:
                    logger.warning(f"PDF file has no pages: {file_path}")
                    return ""
                
                for page_num in range(len(reader.pages)):
                    try:
                        page_text = reader.pages[page_num].extract_text()
                        text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num}: {str(e)}")
                
                return text.strip()
        except PyPDF2.errors.PdfReadError as e:
            logger.error(f"PDF read error for {file_path}: {str(e)}")
            return DocumentParser.fallback_pdf_extraction(file_path)
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def fallback_pdf_extraction(file_path: Path) -> str:
        """
        Fallback method for PDF text extraction when primary method fails.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text
        """
        try:
            logger.info(f"Attempting fallback PDF extraction for {file_path}")
            with open(file_path, 'rb') as file:
                # Use non-strict mode for more permissive parsing
                reader = PyPDF2.PdfReader(file, strict=False)
                text = ""
                for page_num in range(len(reader.pages)):
                    try:
                        text += reader.pages[page_num].extract_text() + "\n"
                    except:
                        # Continue even if individual pages fail
                        pass
                return text.strip()
        except Exception as e:
            logger.error(f"Fallback PDF extraction also failed for {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_path: Path) -> str:
        """
        Extract text from a DOCX file with robust error handling.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text
        """
        try:
            # Validate file is actually a DOCX
            if not DocumentParser.validate_docx(file_path):
                return DocumentParser.fallback_docx_extraction(file_path)
            
            # Parse the document
            doc = docx.Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs]
            
            # Also extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells if cell.text.strip()])
                    if row_text:
                        tables_text.append(row_text)
            
            # Combine paragraphs and tables
            all_text = paragraphs + tables_text
            text = "\n".join([t for t in all_text if t.strip()])
            
            if not text.strip():
                logger.warning(f"No text extracted from DOCX: {file_path}")
            
            return text
        except PackageNotFoundError:
            logger.error(f"DOCX package not found: {file_path}")
            return DocumentParser.fallback_docx_extraction(file_path)
        except zipfile.BadZipFile:
            logger.error(f"File is not a valid DOCX (not a zip file): {file_path}")
            return DocumentParser.fallback_docx_extraction(file_path)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            return DocumentParser.fallback_docx_extraction(file_path)
    
    @staticmethod
    def validate_docx(file_path: Path) -> bool:
        """
        Validate if a file is a proper DOCX file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if file is a valid DOCX, False otherwise
        """
        try:
            # Check if it's a valid ZIP file first
            if not zipfile.is_zipfile(file_path):
                logger.warning(f"File is not a valid ZIP: {file_path}")
                return False
            
            # Check for key DOCX components
            with zipfile.ZipFile(file_path) as zip_ref:
                required_files = ['word/document.xml', '[Content_Types].xml']
                zip_contents = zip_ref.namelist()
                
                for req_file in required_files:
                    if req_file not in zip_contents:
                        logger.warning(f"Missing required DOCX component: {req_file}")
                        return False
            
            return True
        except Exception as e:
            logger.warning(f"DOCX validation failed for {file_path}: {str(e)}")
            return False
    
    @staticmethod
    def fallback_docx_extraction(file_path: Path) -> str:
        """
        Attempt alternative methods to extract text from invalid/corrupted DOCX files.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text or empty string
        """
        logger.info(f"Attempting fallback DOCX extraction for {file_path}")
        
        extracted_text = ""
        
        # Try to extract text from document.xml if it exists
        try:
            if zipfile.is_zipfile(file_path):
                with zipfile.ZipFile(file_path) as zip_ref:
                    if 'word/document.xml' in zip_ref.namelist():
                        with zip_ref.open('word/document.xml') as xml_file:
                            content = xml_file.read().decode('utf-8', errors='ignore')
                            # Basic XML text extraction
                            import re
                            text = re.sub(r'<[^>]+>', ' ', content)
                            text = re.sub(r'\s+', ' ', text).strip()
                            extracted_text = text
                            logger.info(f"Extracted {len(text)} characters from document.xml")
        except Exception as e:
            logger.warning(f"XML extraction failed: {str(e)}")
        
        # If nothing was extracted, try plain text
        if not extracted_text:
            try:
                with open(file_path, 'r', errors='ignore') as f:
                    extracted_text = f.read()
                    logger.info(f"Extracted {len(extracted_text)} characters as plain text")
            except Exception as e:
                logger.warning(f"Plain text extraction failed: {str(e)}")
        
        return extracted_text
    
    @staticmethod
    def extract_text_from_txt(file_path: Path) -> str:
        """
        Extract text from a plain text file with robust error handling.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Extracted text
        """
        try:
            with open(file_path, 'r', errors='ignore') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def auto_detect_and_extract(file_path: Path) -> str:
        """
        Attempt to detect file type and extract text accordingly.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text
        """
        # Check if it's a zip file (potential DOCX)
        if zipfile.is_zipfile(file_path):
            logger.info(f"File appears to be a ZIP archive, attempting DOCX extraction: {file_path}")
            return DocumentParser.extract_text_from_docx(file_path)
        
        # Check if it's a PDF by looking at file signature
        try:
            with open(file_path, 'rb') as f:
                header = f.read(4)
                if header.startswith(b'%PDF'):
                    logger.info(f"File appears to be a PDF based on signature: {file_path}")
                    return DocumentParser.extract_text_from_pdf(file_path)
        except Exception:
            pass
        
        # Try each extraction method in sequence
        methods = [
            DocumentParser.extract_text_from_pdf,
            DocumentParser.extract_text_from_docx,
            DocumentParser.extract_text_from_txt
        ]
        
        for method in methods:
            try:
                text = method(file_path)
                if text.strip():
                    logger.info(f"Successfully extracted text using {method.__name__}")
                    return text
            except Exception:
                pass
        
        logger.error(f"Could not extract text from file {file_path} with any method")
        return ""
    
    @staticmethod
    def from_uploaded_file(file_data: bytes, filename: str) -> str:
        """
        Extract text from an uploaded file (bytes data).
        
        Args:
            file_data: The binary content of the file
            filename: Original filename with extension
            
        Returns:
            Extracted text
        """
        # Save to temporary file and extract
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as temp_file:
            temp_file.write(file_data)
            temp_path = temp_file.name
        
        try:
            text = DocumentParser.extract_text(temp_path)
            return text
        finally:
            # Clean up temporary file
            try:
                os.unlink(temp_path)
            except:
                pass